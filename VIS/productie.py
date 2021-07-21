from flask import (
    Blueprint, flash, g, redirect, render_template, request, url_for
)
from werkzeug.exceptions import abort
from docx import Document
from datetime import date
from docx2pdf import convert
import time

from VIS.auth import login_required
from VIS.db import get_db
from VIS.voorraad import check_prijs, get_afdelingen, get_personeelnummer, get_medewerkers

bp = Blueprint('productie', __name__, url_prefix='/productie')

def change_voornaam(naam):
    voorletter = naam[0]
    voornaam = naam.split()
    achternaam1 = voornaam[1:]
    achternaam2 = ' '
    achternaam = achternaam2.join(achternaam1)
    volledigeNaam1 = voorletter + '.', "(" + voornaam[0] + ")", achternaam
    volledigeNaam2 = ' '
    volledigeNaam = volledigeNaam2.join(volledigeNaam1)
    return volledigeNaam

def create_gebruikersovereenkomst(directieType, naam, functie, artikelType, artikelnaam, merk, serieNummer):
    today = date.today()
    if directieType == 'Algemeen':
        document = Document(r'\\192.168.0.12\Data\ICT\VIS\gebruikersovereenkomsten\Gebruikersovereenkomst2021Algemeen.docx')
    elif directieType == 'Operationeel':
        document = Document(r'\\192.168.0.12\Data\ICT\VIS\gebruikersovereenkomsten\Gebruikersovereenkomst2021Operationeel.docx')
    elif directieType == 'Financieel':
        document = Document(r'\\192.168.0.12\Data\ICT\VIS\gebruikersovereenkomsten\Gebruikersovereenkomst2021Financieel.docx')

    for para in document.paragraphs:
        for run in para.runs:
            if '!naam' in run.text:
                run.text = run.text.replace('!naam', naam)
            elif '!functie' in run.text:
                run.text = run.text.replace('!functie', functie)
            elif '!artikeltype' in run.text:
                run.text = run.text.replace('!artikeltype', artikelType)
            elif 'artikelnaam' in run.text:
                run.text = run.text.replace('artikelnaam', artikelnaam)
            elif '!merk' in run.text:
                run.text = run.text.replace('!merk', merk)
            elif '!serienummer' in run.text:
                run.text = run.text.replace('!serienummer', serieNummer)

    document.save(rf'\\192.168.0.12\Data\ICT\VIS\gebruikersovereenkomsten\{today} - {naam}.docx')
    # time.sleep(5)
    # convert(rf'\\192.168.0.12\Data\ICT\VIS\gebruikersovereenkomsten\{today} - {naam}.docx', rf'\\192.168.0.12\Data\ICT\VIS\gebruikersovereenkomst\{today} - {naam}.pdf')

def get_post(ProductieID):
    ProductieProduct = get_db().execute(
         ' SELECT ProductieID, pro.Artikelnummer, Artikelnaam, Merk, Categorie, Prijs, Opmerking, behA.Gebruikersnaam as GemaaktDoor, '
         ' CAST(CreatieTijd AS smalldatetime) AS CreatieTijd, behB.Gebruikersnaam as UitgifteDoor, CAST(UitgifteTijd AS smalldatetime) AS UitgifteTijd, pro.Personeelnummer, Naam, Afdeling'
         ' FROM Productie pro '
         ' JOIN Artikel art ON pro.Artikelnummer = art.Artikelnummer'
         ' LEFT JOIN Beheerder behA ON pro.GemaaktDoor = behA.GebruikerID'
         ' LEFT JOIN Beheerder behB ON pro.UitgifteDoor = behB.GebruikerID'
		 ' JOIN Medewerker med on pro.Personeelnummer = med.Personeelnummer'
		 ' WHERE ProductieID = ?',
        (ProductieID,)
    ).fetchone()

    if ProductieID is None:
        abort(404, f"Artikelnummer {ProductieID} doesn't exist.")

    return ProductieProduct

@bp.route('/')
def index():
    try:
        productie = get_db().execute(
            ' SELECT ProductieID, pro.Artikelnummer, Artikelnaam, Merk, Categorie, Prijs, Opmerking, behA.Gebruikersnaam as GemaaktDoor, '
            ' CAST(CreatieTijd AS smalldatetime) AS CreatieTijd, behB.Gebruikersnaam as UitgifteDoor, CAST(UitgifteTijd AS smalldatetime) AS UitgifteTijd, pro.Personeelnummer, Naam, Afdeling'
            ' FROM Productie pro '
            ' JOIN Artikel art ON pro.Artikelnummer = art.Artikelnummer'
            ' LEFT JOIN Beheerder behA ON pro.GemaaktDoor = behA.GebruikerID'
            ' LEFT JOIN Beheerder behB ON pro.UitgifteDoor = behB.GebruikerID'
            ' JOIN Medewerker med on pro.Personeelnummer = med.Personeelnummer'
        ).fetchall()
        return render_template('productie/index.html', productie=productie)

    except Exception as e:
        print(e)
        error = 'Er is iets fout gegaan. Je bent teruggestuurd naar de home pagina.'
        flash(error)
        return redirect(url_for('index.index'))

@bp.route('/<int:ProductieID>/update', methods=('GET', 'POST'))
@login_required
def update(ProductieID):
    try:
        if request.method == 'POST':
            prijs = request.form['prijs']
            opmerking = request.form['opmerking']
            naam = request.form['medewerker']
            afdeling = request.form['afdeling']
            error = None

            if prijs.isdecimal() == False:
                prijs, error = check_prijs(prijs)

            if error is not None:
                flash(error)
            else:
                db = get_db()
                db.execute(
                    ' UPDATE Productie SET Prijs  = ? , Opmerking = ? , Personeelnummer = ? '
                    ' WHERE ProductieID = ? ',
                    (prijs, opmerking, get_personeelnummer(naam, afdeling),ProductieID)
                )
                db.commit()
                flash('Item is geüpdatet.')
                return redirect(url_for('productie.index'))
        return render_template('productie/update.html', product=get_post(ProductieID), afdelingen=get_afdelingen(), medewerkers=get_medewerkers())

    except Exception as e:
        print(e)
        error = 'Er is iets fout gegaan. Je bent teruggestuurd naar de home pagina.'
        flash(error)
        return redirect(url_for('index.index'))

@bp.route('/<int:ProductieID>/')
@login_required
def product(ProductieID):
    return render_template('productie/product.html', product=get_post(ProductieID))

@bp.route('/<int:ProductieID>/delete', methods=('POST',))
@login_required
def delete(ProductieID):
    db = get_db()
    print('')
    db.execute(' INSERT INTO HProductie'
               ' SELECT * FROM Productie'
               ' WHERE ProductieID = ?;'
               ' DELETE FROM Productie WHERE ProductieID = ?;', (ProductieID, ProductieID))
    db.commit()
    flash('Item is verwijderd en uit productie')
    return redirect(url_for('voorraad.index'))

@bp.route('/<int:ProductieID>/giveback', methods=('POST', 'GET'))
@login_required
def giveback(ProductieID):
    db = get_db()
    db.execute(' SET IDENTITY_INSERT Voorraad ON;'
               ' INSERT INTO Voorraad (VoorraadID, Artikelnummer, Prijs, GemaaktDoor, CreatieTijd, Opmerking)'
               ' SELECT ProductieID, Artikelnummer, Prijs, GemaaktDoor, CreatieTijd, Opmerking'
               ' FROM Productie '
               ' WHERE ProductieID = ?; '
               ' SET IDENTITY_INSERT Voorraad OFF; '
               ' DELETE FROM Productie WHERE ProductieID = ? ;',
               (ProductieID, ProductieID))
    db.commit()
    flash('Item is uit productie gehaald en terug naar de voorraad verplaatst.')
    return redirect(url_for('voorraad.index'))

@bp.route('/<int:ProductieID>/gebruikersovereenkomst', methods=('POST', 'GET'))
@login_required
def gebruikersovereenkomst(ProductieID):
    try:
        if request.method == 'POST':
            directieType = request.form['directieType']
            naam = request.form['naam']
            functie = request.form['functie']
            categorie = request.form['artikelType']
            artikelnaam = request.form['artikelnaam']
            merk = request.form['merk']
            serieNummer = request.form['serieNummer']

            error = None

            if error is not None:
                flash(error)
            else:
                naam = change_voornaam(naam)
                create_gebruikersovereenkomst(directieType, naam, functie, categorie, artikelnaam, merk, serieNummer)

        return render_template('productie/gebruikersovereenkomst.html', product=get_post(ProductieID))

    except Exception as e:
        print(e)
        error = 'Er is iets fout gegaan. Je bent teruggestuurd naar de home pagina.'
        flash(error)
        return redirect(url_for('index.index'))